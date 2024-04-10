from docx import Document
import re
import os
import streamlit as st


def extract_items(doc_path):
    doc = Document(doc_path)
    items = {}
    current_key = None
    current_content = []

    for para in doc.paragraphs:
        if re.match(r'\d+\.', para.text.strip()):
            if current_key is not None:
                items[current_key] = '\n'.join(current_content).strip()

            current_key = para.text.split('.')[0].strip()
            current_content = [para.text.strip()]
        else:
            current_content.append(para.text.strip())

    if current_key is not None:
        items[current_key] = '\n'.join(current_content).strip()

    return items


def combine_question_answer(questions, answers):
    combined_items = {}
    for key, question in questions.items():
        # Extract just the analysis text without the "1. D 【解析】" part.
        analysis_match = re.search(r'\【解析】(.+)', answers.get(key, ''))
        if analysis_match:
            analysis = analysis_match.group(1).strip()
            # Insert the answer into the question text
            answer_match = re.search(r'^\d+\.\s+([A-E])', answers.get(key, ''))
            if answer_match:
                answer = answer_match.group(1)
                # Assuming there's only one pair of parentheses
                question_with_answer = re.sub(r'\(\s*[A-E]?\s*\)', f'({answer})', question)
                combined_items[key] = f"{question_with_answer}\n解析：{analysis}"
            else:
                combined_items[key] = question + '\n 解析：' + analysis
        else:
            # If the analysis is missing or malformed, keep the original format.
            combined_items[key] = question

    return combined_items


# # Replace with the actual paths
# questions = extract_items('1.docx')
# answers = extract_items('2.docx')
#
# combined_items = combine_question_answer(questions, answers)
#
# combined_doc = Document()
# for key in sorted(combined_items.keys(), key=int):
#     paragraphs = combined_items[key].split('\n')
#     for para in paragraphs:
#         combined_doc.add_paragraph(para)
#
# # Save the combined document.
# combined_doc_path = './combined_document.docx'
# combined_doc.save(combined_doc_path)


# Function to open file dialog and select a file
# def select_file(entry):
#     file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
#     if file_path:
#         entry.delete(0, tk.END)
#         entry.insert(0, file_path)
#
#
# # Function to select directory
# def select_directory(entry):
#     folder_path = filedialog.askdirectory()
#     if folder_path:
#         entry.delete(0, tk.END)
#         entry.insert(0, folder_path)
#
#
# # Function to process the documents
# def process_documents(q_path, a_path, save_path):
#     questions = extract_items(q_path)
#     answers = extract_items(a_path)
#     combined_items = combine_question_answer(questions, answers)
#
#     combined_doc = Document()
#     for key in sorted(combined_items.keys(), key=int):
#         paragraphs = combined_items[key].split('\n')
#         for para in paragraphs:
#             combined_doc.add_paragraph(para)
#
#     combined_doc.save(f"{save_path}/combined_document.docx")
#     messagebox.showinfo("Success", "The documents have been combined successfully!")
#
#
# # Set up the main application window
# root = tk.Tk()
# root.title("Document Combiner")
#
# # Entry fields
# question_entry = tk.Entry(root, width=50)
# question_entry.pack()
#
# answer_entry = tk.Entry(root, width=50)
# answer_entry.pack()
#
# save_entry = tk.Entry(root, width=50)
# save_entry.pack()
#
# # Buttons
# question_button = tk.Button(root, text="选择您的题目文档", command=lambda: select_file(question_entry))
# question_button.pack()
#
# answer_button = tk.Button(root, text="选择您的解析文档", command=lambda: select_file(answer_entry))
# answer_button.pack()
#
# save_button = tk.Button(root, text="选择您保存的地址", command=lambda: select_directory(save_entry))
# save_button.pack()
#
# # Process button
# process_button = tk.Button(root, text="开始合并",
#                            command=lambda: process_documents(question_entry.get(), answer_entry.get(),
#                                                              save_entry.get()))
# process_button.pack()
# copyright_label = Label(root, text="© 2024 ryy. All rights reserved.")
# copyright_label.pack(side=tk.BOTTOM)
# root.mainloop()


# Function to process the documents
def process_documents(q_file, a_file):
    q_path = q_file.name
    a_path = a_file.name

    questions = extract_items(q_path)
    answers = extract_items(a_path)
    combined_items = combine_question_answer(questions, answers)

    combined_doc = Document()
    for key in sorted(combined_items.keys(), key=int):
        paragraphs = combined_items[key].split('\n')
        for para in paragraphs:
            combined_doc.add_paragraph(para)

    # Save the combined document to a temporary file and return the path
    combined_doc_path = 'combined_document.docx'
    combined_doc.save(combined_doc_path)
    return combined_doc_path


# Streamlit app layout
st.title("Document Combiner")

# File uploader widgets
question_file = st.file_uploader("Select the question document", type='docx')
answer_file = st.file_uploader("Select the answer document", type='docx')

# Button to combine documents
if st.button("Combine Documents"):
    if question_file is not None and answer_file is not None:
        # Save uploaded files to disk before processing
        with open(question_file.name, "wb") as f:
            f.write(question_file.getbuffer())
        with open(answer_file.name, "wb") as f:
            f.write(answer_file.getbuffer())

        # Process documents
        combined_path = process_documents(question_file, answer_file)

        # Download button for the combined document
        with open(combined_path, "rb") as f:
            st.download_button(
                label="Download Combined Document",
                data=f,
                file_name="combined_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        # Remove the temporary files
        os.remove(question_file.name)
        os.remove(answer_file.name)
    else:
        st.error("Please upload both question and answer documents.")

# Footer
st.markdown("---")
st.markdown("© 2024 ryy to 20240410. All rights reserved.")
